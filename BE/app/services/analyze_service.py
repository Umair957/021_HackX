import os
from typing import Dict, List, Optional, Tuple
from google import genai
from google.genai import types
from pathlib import Path
from fastapi import HTTPException, status
import json
import re
import uuid
from datetime import datetime
from app.database.models.resume_analysis import (
    ResumeAnalysis, 
    AnalysisScores, 
    ImprovementSuggestion,
    JobContext
)
from app.utils.logger import get_logger

logger = get_logger(__name__)

# Configure Gemini API
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY environment variable is not set")

client = genai.Client(api_key=GEMINI_API_KEY)

class ResumeAnalyzerService:
    """Service for analyzing resumes using Google Gemini AI with Google Search integration"""
    
    def __init__(self):
        # Store client reference
        self.client = client
    
    async def extract_professional_links(self, file_path: str) -> List[str]:
        """Extract professional links (GitHub, LinkedIn, portfolio) from resume using Gemini"""
        try:
            logger.info("Extracting professional links from resume...")
            uploaded_file = self.client.files.upload(file=file_path)
            
            prompt = """Extract all professional links from this resume. Look for:
- GitHub profiles
- LinkedIn profiles  
- Personal websites/portfolios
- Professional social media
- Project links
- Publication links

Return ONLY a JSON array of URLs, like: ["url1", "url2", "url3"]
If no links found, return an empty array: []
"""
            
            response = self.client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[prompt, uploaded_file]
            )
            
            # Delete the uploaded file
            try:
                self.client.files.delete(name=uploaded_file.name)
            except:
                pass
            
            result_text = response.text.strip()
            
            # Clean up markdown
            if result_text.startswith("```json"):
                result_text = result_text[7:]
            if result_text.startswith("```"):
                result_text = result_text[3:]
            if result_text.endswith("```"):
                result_text = result_text[:-3]
            result_text = result_text.strip()
            
            links = json.loads(result_text)
            
            # Filter for professional sites only
            professional_domains = ['github.com', 'linkedin.com', 'gitlab.com', 'bitbucket.org', 
                                   'stackoverflow.com', 'dev.to', 'medium.com']
            filtered_links = [
                link for link in links 
                if any(domain in link.lower() for domain in professional_domains) or 
                re.match(r'https?://[\w\-\.]+\.(io|dev|tech|me|com|net|org)', link.lower())
            ]
            
            logger.info(f"Extracted {len(filtered_links)} professional links")
            return filtered_links[:5]  # Limit to top 5 links
            
        except Exception as e:
            logger.warning(f"Failed to extract links: {str(e)}")
            return []
    
    async def search_candidate_online(self, links: List[str], candidate_name: str = "") -> str:
        """Use Google Search to gather additional information about the candidate"""
        if not links and not candidate_name:
            return ""
        
        try:
            logger.info(f"Searching online for candidate information...")
            
            search_query = f"Analyze and summarize the professional profiles and projects: {' '.join(links)}"
            if candidate_name:
                search_query += f" for {candidate_name}"
            
            response = self.client.models.generate_content(
                model='gemini-2.5-flash',
                contents=search_query,
                config=types.GenerateContentConfig(
                    tools=[types.Tool(google_search=types.GoogleSearch())]
                )
            )
            
            online_info = response.text.strip()
            logger.info(f"Found online information: {len(online_info)} characters")
            return online_info
            
        except Exception as e:
            logger.warning(f"Failed to search online: {str(e)}")
            return ""
    
    def create_analysis_prompt(
        self, 
        job_title: Optional[str] = None, 
        job_description: Optional[str] = None,
        online_info: Optional[str] = None,
        professional_links: Optional[List[str]] = None
    ) -> str:
        """Create a detailed prompt for Gemini to analyze the resume"""
        
        # Base prompt
        prompt = "You are an expert resume analyzer and career coach. Analyze the provided resume file and provide a comprehensive evaluation.\n\n"
        
        # Add online information if available
        if online_info:
            prompt += "**ADDITIONAL ONLINE INFORMATION FOUND:**\n"
            prompt += f"{online_info}\n\n"
            prompt += "**IMPORTANT**: Compare the resume content with the online information. "
            prompt += "Identify any skills, projects, or achievements found online but NOT mentioned in the resume. "
            prompt += "Suggest adding these to strengthen the resume.\n\n"
        
        if professional_links:
            prompt += f"**PROFESSIONAL LINKS FOUND**: {', '.join(professional_links)}\n\n"
        
        # Add job-specific context if provided
        if job_title or job_description:
            prompt += "**TARGET JOB CONTEXT:**\n"
            if job_title:
                prompt += f"- Target Role: {job_title}\n"
            if job_description:
                prompt += f"- Job Description:\n{job_description}\n\n"
            prompt += "Please analyze the resume specifically against this job role and requirements. "
            prompt += "Focus on how well the candidate's experience, skills, and achievements align with the job requirements.\n\n"
        
        prompt += """Please analyze this resume and provide your response in the following JSON format:

{
  "score": <overall score from 0-100>,
  "ats_score": <ATS compatibility score from 0-100>,
  "readability_score": <readability score from 0-100>,
  "keyword_match": <keyword match percentage from 0-100>,
  "strengths": [
    "strength 1",
    "strength 2",
    "strength 3",
    "strength 4"
  ],
  "weaknesses": [
    "weakness 1",
    "weakness 2",
    "weakness 3",
    "weakness 4"
  ],
  "suggestions": [
    {
      "category": "category name",
      "issue": "specific issue",
      "fix": "actionable recommendation",
      "priority": "high|medium|low"
    }
  ]
}

Evaluation Criteria:
1. **Overall Score**: Comprehensive assessment of resume quality (0-100)"""
        
        if job_title or job_description:
            prompt += "\n   - Consider alignment with target job requirements"
        
        prompt += """
2. **ATS Score**: How well the resume will pass Applicant Tracking Systems"""
        
        if job_description:
            prompt += "\n   - Check for keywords from job description"
        
        prompt += """
3. **Readability Score**: How easy it is to read and scan
4. **Keyword Match**: Presence of relevant industry keywords"""
        
        if job_description:
            prompt += "\n   - Match keywords specifically from the provided job description"
        
        prompt += """

Please provide:
- At least 4 strengths (things done well)"""
        
        if job_title:
            prompt += f"\n  - Highlight strengths relevant to {job_title} role"
        
        prompt += """
- At least 4 weaknesses (areas for improvement)"""
        
        if job_description:
            prompt += "\n  - Point out missing skills or experiences from job requirements"
        
        if online_info:
            prompt += "\n  - Mention any impressive achievements found online but missing from resume"
        
        prompt += """
- At least 5 specific, actionable suggestions with priority levels"""
        
        if job_description:
            prompt += "\n  - Prioritize suggestions that improve job description alignment"
        
        if online_info:
            prompt += "\n  - Suggest adding notable online achievements to the resume"
        
        prompt += """
- Be specific and provide actionable feedback
- Focus on content, formatting, ATS compatibility, and impact"""
        
        if online_info:
            prompt += "\n- Highlight any discrepancies between resume and online presence"
        
        prompt += """

Return ONLY the JSON object, no additional text or markdown formatting.
"""
        return prompt
    
    async def analyze_resume(
        self, 
        file_path: str, 
        user_id: str,
        file_name: str,
        file_size: int,
        file_type: str,
        job_title: Optional[str] = None, 
        job_description: Optional[str] = None
    ) -> Dict:
        """
        Main method to analyze a resume file using Gemini AI and save results to database
        
        Args:
            file_path: Path to the resume file
            user_id: ID of the user who owns the resume
            file_name: Original filename
            file_size: File size in bytes
            file_type: MIME type of the file
            job_title: Optional target job title
            job_description: Optional job description for targeted analysis
            
        Returns:
            Dictionary containing analysis results
        """
        uploaded_file = None
        try:
            # 1. Validate file exists
            if not os.path.exists(file_path):
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"File not found: {file_path}"
                )
            
            # 2. Upload file to Gemini
            logger.info(f"Uploading file to Gemini: {Path(file_path).name}")
            if job_title:
                logger.info(f"Target Job Title provided: {job_title[:30]}...")
            if job_description:
                logger.info(f"Job Description provided: {len(job_description)} characters")
            
            uploaded_file = self.client.files.upload(file=file_path)
            logger.info(f"File uploaded successfully to Gemini")
            
            # 3. Extract professional links from resume
            logger.info("Extracting professional links from resume...")
            professional_links = await self.extract_professional_links(file_path)
            
            # 4. Search for additional information online if links found
            online_info = None
            if professional_links:
                logger.info(f"Found {len(professional_links)} professional links, searching online...")
                # Try to extract candidate name from filename or use a generic search
                candidate_name = file_name.replace('.pdf', '').replace('.docx', '').replace('.doc', '').replace('_', ' ')
                online_info = await self.search_candidate_online(professional_links, candidate_name)
            else:
                logger.info("No professional links found in resume")
            
            # 5. Create analysis prompt with job context and online information
            prompt = self.create_analysis_prompt(job_title, job_description, online_info, professional_links)
            
            # 6. Send to Gemini for analysis with the uploaded file
            logger.info("Sending file to Gemini for analysis...")
            response = self.client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[
                    prompt,
                    uploaded_file
                ]
            )
            
            # 7. Parse the response
            result_text = response.text.strip()
            logger.info("Received response from Gemini")
            
            # Remove markdown code blocks if present
            if result_text.startswith("```json"):
                result_text = result_text[7:]
            if result_text.startswith("```"):
                result_text = result_text[3:]
            if result_text.endswith("```"):
                result_text = result_text[:-3]
            
            result_text = result_text.strip()
            
            # 8. Parse JSON response
            analysis_result = json.loads(result_text)
            
            # 9. Validate the structure
            required_keys = ["score", "ats_score", "readability_score", "keyword_match", 
                           "strengths", "weaknesses", "suggestions"]
            
            for key in required_keys:
                if key not in analysis_result:
                    raise ValueError(f"Missing required key: {key}")
            
            # 10. Save to database
            await self._save_to_database(
                user_id=user_id,
                file_name=file_name,
                file_size=file_size,
                file_type=file_type,
                job_title=job_title,
                job_description=job_description,
                analysis_result=analysis_result,
                raw_response=result_text,
                professional_links=professional_links,
                online_info=online_info
            )
            
            # Add links and online info to response
            analysis_result["professional_links"] = professional_links or []
            analysis_result["online_info"] = online_info or None
            
            return analysis_result
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {str(e)}")
            logger.debug(f"Response text length: {len(result_text) if 'result_text' in locals() else 0}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to parse AI response: {str(e)}"
            )
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error analyzing resume: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error analyzing resume: {str(e)}"
            )
        finally:
            # Clean up uploaded file from Gemini if it exists
            if uploaded_file:
                try:
                    self.client.files.delete(name=uploaded_file.name)
                    logger.info("Deleted uploaded file from Gemini")
                except Exception as e:
                    logger.warning(f"Failed to delete file from Gemini: {str(e)}")
    
    async def _save_to_database(
        self,
        user_id: str,
        file_name: str,
        file_size: int,
        file_type: str,
        job_title: Optional[str],
        job_description: Optional[str],
        analysis_result: Dict,
        raw_response: str,
        professional_links: Optional[List[str]] = None,
        online_info: Optional[str] = None
    ) -> None:
        """
        Save analysis results to MongoDB
        
        Args:
            user_id: ID of the user
            file_name: Original filename
            file_size: File size in bytes
            file_type: MIME type
            job_title: Optional job title
            job_description: Optional job description
            analysis_result: Parsed AI analysis result
            raw_response: Raw JSON string from AI
        """
        try:
            # Create JobContext if job details provided
            job_context = None
            if job_title or job_description:
                job_context = JobContext(
                    job_title=job_title,
                    job_description=job_description
                )
            
            # Create AnalysisScores
            scores = AnalysisScores(
                overall_score=float(analysis_result["score"]),
                formatting_score=float(analysis_result["readability_score"]),
                content_quality_score=float(analysis_result["score"]),  # Using overall score
                keyword_optimization_score=float(analysis_result["keyword_match"]),
                ats_compatibility_score=float(analysis_result["ats_score"])
            )
            
            # Convert suggestions to ImprovementSuggestion objects
            suggestions = [
                ImprovementSuggestion(
                    category=s.get("category", "general"),
                    suggestion=f"{s.get('issue', '')}: {s.get('fix', '')}",
                    priority=s.get("priority", "medium")
                )
                for s in analysis_result.get("suggestions", [])
            ]
            
            # Create ResumeAnalysis document
            resume_analysis = ResumeAnalysis(
                user_id=user_id,
                file_name=file_name,
                file_size=file_size,
                file_type=file_type,
                job_context=job_context,
                scores=scores,
                strengths=analysis_result.get("strengths", []),
                weaknesses=analysis_result.get("weaknesses", []),
                improvement_suggestions=suggestions,
                professional_links=professional_links,
                online_info=online_info,
                raw_analysis=raw_response
            )
            
            # Save to database
            await resume_analysis.insert()
            logger.info("Successfully saved analysis to database")
            
        except Exception as e:
            # Log error but don't fail the request
            logger.error(f"Error saving to database: {str(e)}")
            # We don't raise exception here because the analysis was successful
            # Database save is supplementary
    
    @staticmethod
    async def cleanup_temp_file(file_path: str) -> None:
        """Delete temporary file after analysis"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info("Temporary file cleaned up")
        except Exception as e:
            logger.warning(f"Failed to cleanup temp file: {str(e)}")
    
    async def apply_fix_to_content(self, cv_content: str, fix_instruction: str, category: str) -> str:
        """
        Apply an AI suggestion to the CV content and return the modified content.
        
        Args:
            cv_content: The current CV content (plain text)
            fix_instruction: The fix to apply (from AI suggestions)
            category: The category of the fix (e.g., 'Formatting', 'Content')
        
        Returns:
            Modified CV content as a string
        """
        try:
            logger.info(f"Applying fix for category: {category}")
            
            prompt = f"""You are an expert resume writer and editor. You have been given a resume and a specific improvement to make.

**RESUME CONTENT:**
{cv_content}

**CATEGORY:** {category}
**IMPROVEMENT TO APPLY:** {fix_instruction}

**INSTRUCTIONS:**
1. Apply ONLY this specific improvement to the resume
2. Maintain the overall structure and formatting of the resume
3. Keep all other content unchanged unless directly related to this improvement
4. Return the complete modified resume content
5. Do NOT add explanations or comments - return ONLY the modified resume text

Return the complete improved resume:"""

            response = self.client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt
            )
            
            modified_content = response.text.strip()
            
            # Clean up any markdown formatting if present
            if modified_content.startswith("```"):
                lines = modified_content.split("\n")
                modified_content = "\n".join(lines[1:-1]) if len(lines) > 2 else modified_content
                modified_content = modified_content.strip()
            
            logger.info("Successfully applied fix to CV content")
            return modified_content
            
        except Exception as e:
            logger.error(f"Error applying fix: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to apply fix: {str(e)}"
            )
    
    async def save_modified_cv(self, modified_content: str, original_filename: str, 
                              user_id: str, output_dir: Path) -> Tuple[str, str]:
        """
        Save the modified CV content in both the original format and PDF.
        
        Args:
            modified_content: The modified CV content (plain text)
            original_filename: Original filename to determine format
            user_id: User ID for filename generation
            output_dir: Directory to save files
        
        Returns:
            Tuple of (original_format_path, pdf_path)
        """
        try:
            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            base_name = f"modified_resume_{user_id}_{timestamp}_{unique_id}"
            
            # Determine original format
            original_ext = Path(original_filename).suffix.lower()
            
            # Save as text file first (intermediate format)
            text_path = output_dir / f"{base_name}.txt"
            text_path.write_text(modified_content, encoding='utf-8')
            logger.info(f"Saved modified content to: {text_path}")
            
            # Try to convert to original format
            original_format_path = None
            if original_ext in ['.docx', '.doc']:
                try:
                    from docx import Document
                    from docx.shared import Pt
                    
                    # Create DOCX with modified content
                    doc = Document()
                    # Add content with basic formatting
                    for paragraph in modified_content.split('\n\n'):
                        if paragraph.strip():
                            p = doc.add_paragraph(paragraph.strip())
                            # Basic formatting
                            for run in p.runs:
                                run.font.size = Pt(11)
                    
                    docx_path = output_dir / f"{base_name}.docx"
                    doc.save(str(docx_path))
                    original_format_path = str(docx_path)
                    logger.info(f"Saved as DOCX: {docx_path}")
                except ImportError:
                    logger.warning("python-docx not installed. Cannot save as DOCX. Install with: pip install python-docx")
                except Exception as e:
                    logger.error(f"Error creating DOCX: {str(e)}")
            
            # Convert to PDF using reportlab
            pdf_path = None
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfgen import canvas
                from reportlab.lib.units import inch
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                
                pdf_file = output_dir / f"{base_name}.pdf"
                c = canvas.Canvas(str(pdf_file), pagesize=letter)
                width, height = letter
                
                # Set up text formatting
                y_position = height - 1 * inch
                c.setFont("Helvetica", 11)
                
                # Write content to PDF
                for line in modified_content.split('\n'):
                    if y_position < 1 * inch:
                        c.showPage()
                        y_position = height - 1 * inch
                        c.setFont("Helvetica", 11)
                    
                    # Handle long lines
                    if len(line) > 90:
                        words = line.split()
                        current_line = ""
                        for word in words:
                            if len(current_line + " " + word) < 90:
                                current_line += " " + word if current_line else word
                            else:
                                c.drawString(1 * inch, y_position, current_line)
                                y_position -= 15
                                current_line = word
                                if y_position < 1 * inch:
                                    c.showPage()
                                    y_position = height - 1 * inch
                                    c.setFont("Helvetica", 11)
                        if current_line:
                            c.drawString(1 * inch, y_position, current_line)
                            y_position -= 15
                    else:
                        c.drawString(1 * inch, y_position, line)
                        y_position -= 15
                
                c.save()
                pdf_path = str(pdf_file)
                logger.info(f"Saved as PDF: {pdf_file}")
            except ImportError:
                logger.warning("reportlab not installed. Cannot save as PDF. Install with: pip install reportlab")
            except Exception as e:
                logger.error(f"Error creating PDF: {str(e)}")
            
            # Return paths (use text file as fallback if conversions failed)
            return (
                original_format_path or str(text_path),
                pdf_path or str(text_path)
            )
            
        except Exception as e:
            logger.error(f"Error saving modified CV: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to save modified CV: {str(e)}"
            )
    
    async def save_modified_cv_with_structure(self, modified_content: str, original_filename: str,
                                             original_file_path: Optional[str], user_id: str, 
                                             output_dir: Path) -> Tuple[str, str]:
        """
        Save the modified CV content while preserving the original document structure.
        
        Args:
            modified_content: The modified CV content (plain text)
            original_filename: Original filename to determine format
            original_file_path: Path to the original file (if available) to preserve structure
            user_id: User ID for filename generation
            output_dir: Directory to save files
        
        Returns:
            Tuple of (original_format_path, pdf_path)
        """
        try:
            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            base_name = f"modified_resume_{user_id}_{timestamp}_{unique_id}"
            
            # Determine original format
            original_ext = Path(original_filename).suffix.lower()
            
            original_format_path = None
            pdf_path = None
            
            # If original file provided and is DOCX, preserve its structure
            if original_file_path and original_ext == '.docx':
                try:
                    from docx import Document
                    from docx.shared import Pt, RGBColor
                    
                    logger.info("Preserving original DOCX structure...")
                    
                    # Load the original document to preserve styles
                    doc = Document(original_file_path)
                    
                    # Clear existing content but keep styles
                    for element in doc.element.body:
                        doc.element.body.remove(element)
                    
                    # Parse modified content and apply to document
                    paragraphs = modified_content.split('\n\n')
                    
                    for para_text in paragraphs:
                        if not para_text.strip():
                            continue
                            
                        # Add paragraph with original document's default style
                        p = doc.add_paragraph(para_text.strip())
                        
                        # Try to detect and apply basic formatting hints from text
                        # Headers (ALL CAPS or starts with common header keywords)
                        if para_text.strip().isupper() or any(para_text.strip().startswith(kw) for kw in ['EXPERIENCE', 'EDUCATION', 'SKILLS', 'SUMMARY', 'CONTACT']):
                            p.style = 'Heading 1' if 'Heading 1' in [s.name for s in doc.styles] else 'Normal'
                            for run in p.runs:
                                run.bold = True
                                run.font.size = Pt(14)
                        # Bullet points
                        elif para_text.strip().startswith(('•', '-', '*', '▪')):
                            p.style = 'List Bullet' if 'List Bullet' in [s.name for s in doc.styles] else 'Normal'
                        # Regular content
                        else:
                            for run in p.runs:
                                run.font.size = Pt(11)
                    
                    docx_path = output_dir / f"{base_name}.docx"
                    doc.save(str(docx_path))
                    original_format_path = str(docx_path)
                    logger.info(f"Saved DOCX with preserved structure: {docx_path}")
                    
                except ImportError:
                    logger.warning("python-docx not installed. Cannot preserve DOCX structure.")
                except Exception as e:
                    logger.error(f"Error preserving DOCX structure: {str(e)}")
            
            # If DOCX creation failed or wasn't applicable, create new document
            if not original_format_path and original_ext in ['.docx', '.doc']:
                try:
                    from docx import Document
                    from docx.shared import Pt, RGBColor
                    
                    doc = Document()
                    
                    # Add content with smart formatting
                    for para_text in modified_content.split('\n\n'):
                        if not para_text.strip():
                            continue
                            
                        p = doc.add_paragraph(para_text.strip())
                        
                        # Apply formatting based on content analysis
                        if para_text.strip().isupper() or any(para_text.strip().startswith(kw) for kw in ['EXPERIENCE', 'EDUCATION', 'SKILLS', 'SUMMARY', 'CONTACT']):
                            for run in p.runs:
                                run.bold = True
                                run.font.size = Pt(14)
                        elif para_text.strip().startswith(('•', '-', '*', '▪')):
                            p.style = 'List Bullet'
                        else:
                            for run in p.runs:
                                run.font.size = Pt(11)
                    
                    docx_path = output_dir / f"{base_name}.docx"
                    doc.save(str(docx_path))
                    original_format_path = str(docx_path)
                    logger.info(f"Saved new DOCX: {docx_path}")
                    
                except ImportError:
                    logger.warning("python-docx not installed.")
                except Exception as e:
                    logger.error(f"Error creating DOCX: {str(e)}")
            
            # Convert to PDF
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfgen import canvas
                from reportlab.lib.units import inch
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.enums import TA_LEFT, TA_CENTER
                
                pdf_file = output_dir / f"{base_name}.pdf"
                doc = SimpleDocTemplate(str(pdf_file), pagesize=letter)
                story = []
                styles = getSampleStyleSheet()
                
                # Create custom styles
                header_style = ParagraphStyle(
                    'CustomHeader',
                    parent=styles['Heading1'],
                    fontSize=14,
                    textColor='#1a1a1a',
                    spaceAfter=12,
                    bold=True
                )
                
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor='#333333',
                    spaceAfter=8
                )
                
                # Parse and add content
                for para_text in modified_content.split('\n\n'):
                    if not para_text.strip():
                        continue
                    
                    # Detect headers
                    if para_text.strip().isupper() or any(para_text.strip().startswith(kw) for kw in ['EXPERIENCE', 'EDUCATION', 'SKILLS', 'SUMMARY', 'CONTACT']):
                        story.append(Paragraph(para_text.strip(), header_style))
                    else:
                        story.append(Paragraph(para_text.strip(), normal_style))
                    
                    story.append(Spacer(1, 0.1 * inch))
                
                doc.build(story)
                pdf_path = str(pdf_file)
                logger.info(f"Saved PDF with formatting: {pdf_file}")
                
            except ImportError:
                logger.warning("reportlab not fully installed. Using basic PDF generation.")
                # Fallback to basic PDF
                try:
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib.units import inch
                    
                    pdf_file = output_dir / f"{base_name}.pdf"
                    c = canvas.Canvas(str(pdf_file), pagesize=letter)
                    width, height = letter
                    y_position = height - 1 * inch
                    c.setFont("Helvetica", 11)
                    
                    for line in modified_content.split('\n'):
                        if y_position < 1 * inch:
                            c.showPage()
                            y_position = height - 1 * inch
                            c.setFont("Helvetica", 11)
                        
                        if len(line) > 90:
                            words = line.split()
                            current_line = ""
                            for word in words:
                                if len(current_line + " " + word) < 90:
                                    current_line += " " + word if current_line else word
                                else:
                                    c.drawString(1 * inch, y_position, current_line)
                                    y_position -= 15
                                    current_line = word
                                    if y_position < 1 * inch:
                                        c.showPage()
                                        y_position = height - 1 * inch
                                        c.setFont("Helvetica", 11)
                            if current_line:
                                c.drawString(1 * inch, y_position, current_line)
                                y_position -= 15
                        else:
                            c.drawString(1 * inch, y_position, line)
                            y_position -= 15
                    
                    c.save()
                    pdf_path = str(pdf_file)
                    logger.info(f"Saved basic PDF: {pdf_file}")
                except Exception as e:
                    logger.error(f"Error creating PDF: {str(e)}")
            except Exception as e:
                logger.error(f"Error creating formatted PDF: {str(e)}")
            
            # Fallback to text file if all conversions failed
            if not original_format_path and not pdf_path:
                text_path = output_dir / f"{base_name}.txt"
                text_path.write_text(modified_content, encoding='utf-8')
                logger.info(f"Saved as text file (fallback): {text_path}")
                return (str(text_path), str(text_path))
            
            # Return paths (use text as fallback if one format failed)
            if not original_format_path:
                text_path = output_dir / f"{base_name}.txt"
                text_path.write_text(modified_content, encoding='utf-8')
                original_format_path = str(text_path)
            
            if not pdf_path:
                pdf_path = original_format_path
            
            return (original_format_path, pdf_path)
            
        except Exception as e:
            logger.error(f"Error saving modified CV with structure: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to save modified CV: {str(e)}"
            )
